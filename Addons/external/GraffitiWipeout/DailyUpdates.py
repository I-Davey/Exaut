from array import array
from asyncio.log import logger
##imports for exaut pyinstaller
from html.parser import HTMLParser

import inspect
from pathlib import Path
from re import S
from __important.PluginInterface import PluginInterface, Types
import os 
from PyQt6.QtWidgets import *
from PyQt6.QtCore import *
from functools import partial
import random
import time

import openpyxl
import pyperclip
import datetime
from bs4 import BeautifulSoup
import json
import requests
import what3words
from docx import Document, table, oxml, opc
from docx.enum.dml import MSO_THEME_COLOR_INDEX
import copy
from .__db import gw_db
from .__orm import *
from sqlalchemy import insert, and_, or_, not_, func


class DailyUpdates(PluginInterface):
    load = True
    types = Types.source | Types.type_ | Types.keyfile | Types.keypath | Types.databasename
   
    type_types = { "__Name":"Daily Updates GraffitiWipeout"}
    callname = ("GWDUpdate", "GWDUpdateImport")
    hooks_handler = ["log"]



    def load_self(self, hooks):
        self.logger = hooks["log"]
        return True

    def main(self, save_path, type_, w3w_api_key, datapase_path, template_file) -> bool: 
        self.w3w_geo = what3words.Geocoder(w3w_api_key)
        self.template_file = template_file
        self.db_loc = datapase_path
        #check if dir exists
        self.save_path = save_path
        if type_ == "GWDUpdateImport":
            dialog = None
            x = self.Popups.custom(dialog)
            print(x)


        #if exe exists:   
        dialog = Dialog
        x = self.Popups.custom(dialog, self.excel_handle_import, self.load_av_session)
        if len(x) == 1 and not x[0]:
            return False
        completed_jobs_data,self.filename = x

        main_array = completed_jobs_data
        
        #dump to excel
        openpxl = openpyxl.Workbook()
        for sheet in main_array:
            ws = openpxl.create_sheet(sheet)
            for row in main_array[sheet]:
                #clean row to be fit as a valid column name

                ws.append(row)
        while True:
            iint = 0
            i = ""
            try:
                openpxl.save(self.save_path + f"/{self.filename}_backup{i}.xlsx")
                break
            except PermissionError:
                if iint > 199:
                    break
                iint += 1
                i = f"_{iint}"


            


        #for each row in main_array, if w3w is a key and value is not empty, convert to gmap link and append pindrop column
        for sheet in main_array.keys():
            header = main_array[sheet][0]
            if "W3W" in header:
                w3w_index = header.index("W3W")
                for i in range(len(main_array[sheet])):
                    if i == 0:
                        main_array[sheet][i].append("pindrop")
                        continue
                    if main_array[sheet][i][w3w_index] != None:
                        gmap_link, main_array = self.handle_gmap_pindrop(main_array[sheet][i][w3w_index], [main_array, sheet, i, w3w_index])
                        main_array[sheet][i].append(gmap_link)
                    else:
                        main_array[sheet][i].append(None)
        self.total_savings(main_array["Completed Jobs"])
        self.json_data = main_array
        self.db_main()

        self.table_main()

        self.excel_handle_export()
        self.Popups.alert("Exported to " + save_path + self.filename)


    def total_savings(self, completed_jobs_data):
        unit_cost = 1583.33
        dedicated_clean_sqm_price = 8
        dedicated_paint_sqm_price = 1
        ad_hoc_clean_sqm_price = 15.25
        ad_hoc_paint_sqm_price = 6.2
        cleaned_sqm = 0
        painted_sqm = 0
        for row in completed_jobs_data:
            if row[0] == "AV#":
                continue
            #remove .0 from sqm
            if type(row[4]) == str:
                row[4] = row[4].split(".")[0] if "." in row[4] else row[4]
            if type(row[3]) == str:
                row[3] = row[3].split(".")[0] if "." in row[3] else row[3]
            cleaned_sqm += int(row[4])
            painted_sqm += int(row[3])

        total_cleaning_cost = cleaned_sqm * dedicated_clean_sqm_price
        total_painting_cost = painted_sqm * dedicated_paint_sqm_price
        total_cost = total_cleaning_cost + total_painting_cost + unit_cost
        
        ad_hoc_clean_cost = ad_hoc_clean_sqm_price * cleaned_sqm
        ad_hoc_paint_cost = ad_hoc_paint_sqm_price * painted_sqm
        ad_hoc_total_cost = ad_hoc_clean_cost + ad_hoc_paint_cost



        total_savings = ad_hoc_total_cost - total_cost
        self.savings =  {"sqp": cleaned_sqm, "sqm": painted_sqm, "total_savings": total_savings}

    def load_av_session(self):
        url = "https://vicroads.assetvision.com.au/api/Authenticate/UserLogin"
        payload = {"userName":"***REMOVED***", "password":"***REMOVED***!"}
        headers = {'Content-Type': 'application/json'}
        #session:
        session = requests.Session()
        response = session.post(url, data=json.dumps(payload), headers=headers)
        next_req = response.json()["SystemLoginSites"][0]["URL"]
        print(next_req)
        next_res = session.get(next_req)
        soup = BeautifulSoup(next_res.content)
        payload["___VIEWSTATE"] = soup.select_one("#__VIEWSTATE")["value"]
        payload["___VIEWSTATE"] = payload["___VIEWSTATE"].replace("+", "%2B").replace("/", "%2F").replace("=", "%3D")
        payload["___VIEWSTATEGENERATOR"] = soup.select_one("#__VIEWSTATEGENERATOR")["value"]
        payload["___VIEWSTATEGENERATOR"] = payload["___VIEWSTATEGENERATOR"].replace("+", "%2B").replace("/", "%2F").replace("=", "%3D")
        payload["___EVENTVALIDATION"] = soup.select_one("#__EVENTVALIDATION")["value"]
        payload["___EVENTVALIDATION"] = payload["___EVENTVALIDATION"].replace("+", "%2B").replace("/", "%2F").replace("=", "%3D")
        payload["___EVENTVALIDATION"] = payload["___EVENTVALIDATION"].replace("+", "%2B").replace("/", "%2F").replace("=", "%3D")
        next_req += "&ReturnUrl=%2f"
        payload_new_urlencoded = f"__EVENTTARGET=&__EVENTARGUMENT=&__VIEWSTATE={payload['___VIEWSTATE']}&__VIEWSTATEGENERATOR={payload['___VIEWSTATEGENERATOR']}&__EVENTVALIDATION={payload['___EVENTVALIDATION']}&hidTimezoneOffset=-660&hidTimezone=Australia%2FSydney&btnLogin="
        session.post(next_req, data=payload_new_urlencoded, headers={'Content-Type': 'application/x-www-form-urlencoded'})
        return session

    def handle_gmap_pindrop(self, w3w, idx):
        if "," in w3w:
            lat, lng = w3w.split(",")
            w3w = self.w3w_geo.convert_to_3wa(what3words.Coordinates(lat, lng))
            idx[0][idx[1]][idx[2]][idx[3]] = w3w["words"]
            return f"https://www.google.com/maps/search/?api=1&query={lat},{lng}", idx[0]
        #elif 2 dots in w3w:
        elif w3w.count(".") == 2:
            coords = self.w3w_geo.convert_to_coordinates(w3w)
            if 'error' in coords:
                return None, idx[0]
            return f"https://www.google.com/maps/search/?api=1&query={coords['coordinates']['lat']},{coords['coordinates']['lng']}", idx[0]
        return None, idx[0]

    def excel_handle_export(self):
        #create excel file
        wb = openpyxl.Workbook()

        wb.remove(wb.active)

        #add data to excel file


        #put each k v as seperate sheet in excel file
        for sheet in self.json_client.keys():
            ws = wb.create_sheet(sheet)
            for row in self.json_client[sheet]:
                #if type list in row
                ws.append(row)

        #delete default sheet
        save_path = self.save_path + "\\" + self.filename + ".xlsx" if not self.filename.endswith(".xlsx") else self.save_path + "\\" + self.filename
        while True:
            try:
                wb.save(save_path)
                break
            except PermissionError:
                    x = self.Popups.yesno("File is open in another program. Please close it and press OK to continue.")
                    if not x:
                        break
            
        self.logger.success("Exported to " + self.save_path + "\\" + self.filename)

    def excel_handle_import(self, file):
        #open excel file
        wb = openpyxl.load_workbook(file)

        #get all sheets
        sheets = wb.sheetnames

        #get all data from sheets
        data = {}
        for sheet in sheets:
            ws = wb[sheet]
            data[sheet] = []
            #all data except first row
            for row in ws.iter_rows(min_row=2):
                data[sheet].append([cell.value for cell in row])

        #return data
        return data


    ##word doc stuff
    def replace_text(self, document, find, replace):
        for p in document.paragraphs:
            if find in p.text:
                inline = p.runs
                # Replace strings
                for i in range(len(inline)):
                    if find in inline[i].text:
                        text = inline[i].text.replace(find, replace)
                        inline[i].text = text

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    self.replace_text(cell, find, replace)

    def get_tables(self, document):
        tables = []
        for table in document.tables:
            tables.append(table)
        return tables

    def get_table_data(self, table):
        data = []
        for row in table.rows:
            data.append([])
            for cell in row.cells:
                data[-1].append(cell.text)
        return data

    def get_table_name(self, table):
        return table.rows[1].cells[0].text

    def add_table_data_noheader(self, table, data):
        #remove all but header row from table
        for i in range(len(table.rows)-1, 0, -1):
            tr = table.rows[i]
            table._tbl.remove(tr._tr)
        #add data
        for row in data:
            row_ = table.add_row().cells
            for i in range(len(row)):
                if type(row[i]) == list and row[i][0] == "URL":

                    self.add_hyperlink(row_[i].add_paragraph(), row[i][1], row[i][2])
                else:
                    row_[i].text = str(row[i])
                
    def add_hyperlink(self, paragraph, text, url):
        # This gets access to the document.xml.rels file and gets a new relation id value
        part = paragraph.part
        r_id = part.relate_to(url, opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

        # Create the w:hyperlink tag and add needed values
        hyperlink = oxml.shared.OxmlElement('w:hyperlink')
        hyperlink.set(oxml.shared.qn('r:id'), r_id, )

        # Create a w:r element and a new w:rPr element
        new_run = oxml.shared.OxmlElement('w:r')
        rPr = oxml.shared.OxmlElement('w:rPr')

        # Join all the xml elements together add add the required text to the w:r element
        new_run.append(rPr)
        new_run.text = text
        hyperlink.append(new_run)

        # Create a new Run object and add the hyperlink into it
        r = paragraph.add_run ()
        r._r.append (hyperlink)

        # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
        # Delete this if using a template that has the hyperlink style in it
        r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
        r.font.underline = True


        return hyperlink

    def table_main(self):
        json_data = self.json_data.copy()
        #replace W3W with google maps link
        #create complete detached copy of json_data
        json_data_xcel = copy.deepcopy(json_data)
        for sheet in json_data.keys():
            header = json_data[sheet][0]
            if "W3W" in header and "pindrop" in header:
                pindrop_loc = header.index("pindrop")
                w3w_loc = header.index("W3W")
                #replace all w3w  with pindrop
                for i in range(1, len(json_data[sheet])):
                    pd_data = json_data[sheet][i][pindrop_loc]
                    #turn pd_data into a docx hyperlink with text "Pindrop"
                    if pd_data:
                        pd_data = ["URL", "Pindrop", pd_data]
                        json_data_xcel[sheet][i][w3w_loc] ="=HYPERLINK(\"{}\", \"{}\")".format(pd_data[2], pd_data[1])
                    else:
                        json_data_xcel[sheet][i][w3w_loc] = ""

                    json_data[sheet][i][w3w_loc] = pd_data
                #rename w3w to pindrop
                header[w3w_loc] = "pindrop"
                #delete pindrop
                del header[pindrop_loc]
                #delete pindrop column
                for i in range(1, len(json_data[sheet])):
                    del json_data[sheet][i][pindrop_loc]
                    del json_data_xcel[sheet][i][pindrop_loc]

                if "Location" in header:
                    loc_loc = header.index("Location")
                    #swap location and pindrop
                    for i in range(1, len(json_data[sheet])):
                        json_data[sheet][i][loc_loc], json_data[sheet][i][w3w_loc] = json_data[sheet][i][w3w_loc], json_data[sheet][i][loc_loc]
                        json_data_xcel[sheet][i][loc_loc], json_data_xcel[sheet][i][w3w_loc] = json_data_xcel[sheet][i][w3w_loc], json_data_xcel[sheet][i][loc_loc]
                
        self.json_client = json_data_xcel
        self.json_db = json_data
        document = Document(self.template_file)
        self.replace_text(document, "{{DATE}}", datetime.datetime.now().strftime("%d/%m/%Y"))
        sqp = self.savings["sqp"]
        sqm = self.savings["sqm"]
        total_savings = self.savings["total_savings"]
        self.replace_text(document, "squarepainted", str(sqp))
        self.replace_text(document, "squarecleaned", str(sqm))
        self.replace_text(document, "totalsaving", str(total_savings))

        tables = self.get_tables(document)
        found_tables = {}
        for table_ in tables:
            print(self.get_table_name(table_))
            if self.get_table_name(table_) == "Completed_jobs":
                found_tables["Completed Jobs"] = table_
            elif self.get_table_name(table_) == "inspections":
                found_tables["Inspection Routes"] = table_
            elif self.get_table_name(table_) == "inspected":
                found_tables["Inspections"] = table_
            elif self.get_table_name(table_) == "schedule":
                found_tables["Schedule"] = table_
            elif self.get_table_name(table_) == "comments":
                #delete first row
                for i in (0,1):
                    tr = table_.rows[0]
                    table_._tbl.remove(tr._tr)
                found_tables["Comments"] = table_
        
        for table in found_tables.keys():
            self.add_table_data_noheader(found_tables[table], json_data[table].copy()[1:])
        while True:
            try:
                document.save(self.save_path + "\\" + self.filename + ".docx" if not self.filename.endswith(".docx") else self.save_path + "\\" + self.filename)
                break
            except PermissionError:
                    x = self.Popups.yesno("File is open in another program. Please close it and press OK to continue.")
                    if not x:
                        break
                

#db stuff
    def db_main(self):
        db_loc = Path(self.db_loc)
        db_con = gw_db(db_loc)
        db_map = {
            "Completed Jobs" : completed_jobs,
            "Inspection Routes" : inspection_routes,
            "Inspections" : inspections,
            "Schedule" : schedule,
            "Comments" : comments
        }
        record_map = {
            "AV#" : "assetvision",
            "W3W" : "w3w",
            "Location" : "location",
            'Sq/m Painted' : "sqm_painted",
            'Sq/m Cleaned' : "sqm_cleaned",
            'TM' : "tm",
            'Notes': "notes",
            'inspection route' : "inspection_route",
            'start': "start",
            'end' : "end",
            "Comments" : "comments"



        }
        for table in self.json_data.keys():
            table_object = db_map[table]
            header = self.json_data[table][0]

            header_map = { header : record_map[header] for header in header if header in record_map.keys()}
           
            #if item not in header_map, it is not in the db
            records = []
            for row in self.json_data[table][1:]:
                #define each row in header as a key

                record = {}
                for i,v in enumerate(row):
                    if header[i] in header_map.keys():
                        record[header_map[header[i]]] = row[i]
                    #TODO: Add Date properly
                record["date"] = datetime.datetime.now().strftime("%d/%m/%Y")
                records.append(insert(table_object).values(record))
            db_con.writesql(records)
            

#################################################
#####XLSX FILE INSERTION WITH PYQT###############
class SelectxlsxFile(QDialog):
    signal = pyqtSignal(str)
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Select xlsx File")
        self.setFixedSize(400, 400)
        self.setStyleSheet("background-color: #1e1e1e; color: #ffffff;")

        self.layout = QVBoxLayout()
        self.setLayout(self.layout)

        self.label = QLabel("Select xlsx File")
        self.layout.addWidget(self.label)

        self.button = QPushButton("Select")
        self.button.clicked.connect(self.select_file)
        self.layout.addWidget(self.button)

        self.file_path = None

    def select_file(self):
        self.file_path = QFileDialog.getOpenFileName(self, "Select Data File", "", "xlsx Files (*.xlsx)")[0]
        self.signal.emit(self.file_path)
        self.close()


#################################################
#####CompletedJobs Handler Object################
class CompletedJobs:
    def __init__(self, parent, session, columns):
        super().__init__()
        self.widget = QWidget(parent)
        self.layout = QVBoxLayout(self.widget)
        self.layout.setAlignment(Qt.AlignmentFlag.AlignTop)
        self.session = session
        self.s = session()
        
        self.completed_jobs_text = QLabel("Completed Jobs")
        self.layout.addWidget(self.completed_jobs_text)
        self.columns = columns
        self.table = QTableWidget()
        self.table.setColumnCount(len(self.columns))
        self.table.setHorizontalHeaderLabels(self.columns)
        #make table add rows to bottom
        #vertical header stretchy
        self.table.verticalHeader().setSectionResizeMode(QHeaderView.ResizeMode.ResizeToContents)
        

        self.table.setRowCount(0)
        self.layout.addWidget(self.table)



        self.av_btn = QPushButton("Load AV")
        self.av_btn.clicked.connect(self.load_av)


        #create a + button for completed_jobs to add a new row
        self.add_btn = QPushButton("+")
        self.add_btn.clicked.connect(partial(self.addrow, self.table))



        self.add_btn.click()
        #create a - button for completed_jobs to remove a row
        self.rm_btn = QPushButton("-")
        self.rm_btn.clicked.connect(partial(self.removerow, self.table))

        plusminushorzntl = QHBoxLayout()
        plusminushorzntl.addWidget(self.av_btn)
        plusminushorzntl.addWidget(self.add_btn)
        plusminushorzntl.addWidget(self.rm_btn)

        self.layout.addLayout(plusminushorzntl)

    def load_av(self):
        #qdialog enter av number
        av_num = QInputDialog.getText(self.widget, "AV Number", "Enter AV Number")[0]
        if not av_num:
            return
        else:
            av_num = int(av_num)
        url = "https://vicroads.assetvision.com.au/Maintenance/LoadJobDetails"
        payload = {"jobId": av_num}
        headers = {'Content-Type': 'application/json'}

        response = self.s.post(url, data=json.dumps(payload), headers=headers)
        if response.url == "https://vicroads.assetvision.com.au/LoginPage/Login.aspx?ReturnUrl=%2fMaintenance%2fLoadJobDetails":
            self.reload_af()
            response = self.s.post(url, data=json.dumps(payload), headers=headers)
        if response.url == "https://vicroads.assetvision.com.au/LoginPage/Login.aspx?ReturnUrl=%2fMaintenance%2fLoadJobDetails":
            #qt alert
            QMessageBox.about(self.widget, "Error", "Login Failed")
            return
        
        data = json.loads(response.json()["StringResult"])
        gpslat = data["GPSLat"]
        gpslong = data["GPSLong"]
        comment = data["AllComments"][0]["Comment"] if len(data["AllComments"]) > 0 else ""
        
        location = data["ContractAssetAndContractRoad"]["FullAssetName"]
        sqm = data["Quantity"]
        #generate w3w from gps
        w3w = f"{gpslat},{gpslong}"
        #add to table
        #count rows in table
        row = self.table.rowCount()
        self.table.insertRow(row)
        self.table.setItem(row, 0, QTableWidgetItem(str(av_num)))
        self.table.setItem(row, 1, QTableWidgetItem(w3w))
        self.table.setItem(row, 2, QTableWidgetItem(location))
        self.table.setItem(row, 3, QTableWidgetItem(str(sqm)))
        self.table.setItem(row, 4, QTableWidgetItem("0"))
        self.table.setItem(row, 5, QTableWidgetItem(comment))

    def reload_af(self):
        self.s = self.session()

    def addrow(self, table, row=None):
        if not row:
            #add row to bottom
            table.setRowCount(table.rowCount() + 1)
        else:
            table.insertRow(row)
        self.setrowtypes()
        #resize the table to fit the new row
    def removerow(self, table, row=None):
        if not row:
            table.setRowCount(table.rowCount() - 1)
        else:
            table.removeRow(row)
        table.resizeRowsToContents()

    def setrowtypes(self):
        row = self.table.rowCount()-1
        sb1 = QSpinBox()
        sb1.setMaximum(2147483647)
        sb2 = QSpinBox()
        sb2.setMaximum(2147483647)
        sb3 = QSpinBox()
        sb3.setMaximum(2147483647)
        self.table.setCellWidget(row,0, sb1)
        #width for column 1 half
        self.table.setColumnWidth(0, 100)
    
        ##column 1 and 2 are stretchy
        self.table.horizontalHeader().setSectionResizeMode(1, QHeaderView.ResizeMode.ResizeToContents)
        self.table.horizontalHeader().setSectionResizeMode(2, QHeaderView.ResizeMode.Stretch)

        self.table.setItem(row, 1, QTableWidgetItem()) # world location
        self.table.setItem(row, 2, QTableWidgetItem()) # World3Word location


        self.table.setCellWidget(row,3,sb2)
        self.table.setColumnWidth(3, 86)

        self.table.setCellWidget(row,4,sb3)
        self.table.setColumnWidth(4, 86)



#################################################
#####SHEET CREATOR OBJECT########################
class TabSheets: 
    def __init__(self, parent=None, name = None, columns = None):
        super().__init__()
        self.widget = QWidget(parent)
        self.layout = QVBoxLayout(self.widget)
        self.layout.setAlignment(Qt.AlignmentFlag.AlignTop)

        self.text = QLabel(name)
        self.layout.addWidget(self.text)

        self.table = QTableWidget()
        self.columns = columns
        self.table.setColumnCount(len(self.columns))
        self.table.setHorizontalHeaderLabels(self.columns.keys())
        self.table.setRowCount(1)
        self.table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch)
        self.layout.addWidget(self.table)
    
        #create a + button for comments to add a new row
        self.add_btn = QPushButton("+")
        self.add_btn.clicked.connect(partial(self.addrow, self.table))
        #create a - button for comments to remove a row
        self.rm_btn = QPushButton("-")
        self.rm_btn.clicked.connect(partial(self.removerow, self.table))

        plusminushorzntl = QHBoxLayout()
        plusminushorzntl.addWidget(self.add_btn)
        plusminushorzntl.addWidget(self.rm_btn)
        self.layout.addLayout(plusminushorzntl)

    def addrow(self, table, row=None):
        if not row:
            #add row to bottom
            table.setRowCount(table.rowCount() + 1)
        else:
            table.insertRow(row)
        self.setrowtypes()
        #resize the table to fit the new row

    def removerow(self, table, row=None):
        if not row:
            table.setRowCount(table.rowCount() - 1)
        else:
            table.removeRow(row)
        table.resizeRowsToContents()

    def setrowtypes(self):
        for col, (colname, coltype) in enumerate(self.columns.items()):
            #int float str bool
            if coltype == "int":
                sb = QSpinBox()
                sb.setMaximum(2147483647)
                self.table.setCellWidget(self.table.rowCount()-1,col, sb)
            elif coltype == "float":
                sb = QDoubleSpinBox()
                sb.setMaximum(2147483647)
                self.table.setCellWidget(self.table.rowCount()-1,col, sb)
            elif coltype == "str":
                self.table.setItem(self.table.rowCount()-1,col, QTableWidgetItem())
            elif coltype == "bool":
                cb = QCheckBox()
                self.table.setCellWidget(self.table.rowCount()-1,col, cb)
                #fit the checkbox to the cell
                self.table.setColumnWidth(col, 20)
                
            else:
                print(f"Unknown column type {coltype} for column {colname}")

#################################################
#####MAIN DIALOG OBJECT@@@@@@@@@@@###############        
class Dialog(QDialog):
    NumGridRows = 3
    NumButtons = 4
    signal = pyqtSignal(tuple)
    _done = False

    def __init__(self, parent, handle_import, av_session, tables = None, data=None):
        super(Dialog, self).__init__(parent)
        self.session = av_session
        self.excel_handle_import = handle_import
        self.tables = tables
        self.data = data
        self.createFormGroupBox()
        
        buttonBox = QDialogButtonBox(QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel)
        buttonBox.accepted.connect(self.accept)
        buttonBox.rejected.connect(self.reject)
        
        mainLayout = QVBoxLayout()
        tb = self.handleToolbar()
        mainLayout.setMenuBar(tb)
        mainLayout.addWidget(self.formGroupBox)
        #current datetime
        h1 = QHBoxLayout()
        h1.addWidget(QLabel("File Name:"))

        self.filename= QLineEdit(f"DailyReport-{datetime.datetime.now().strftime('%Y-%m-%d')}")
        h1.addWidget(self.filename)
        h1.addWidget(buttonBox)
        mainLayout.addLayout(h1)


        self.setLayout(mainLayout)
        
        #increase width to 900, allow you to change the width of the window
        self.resize(900, 300)





        self.setWindowTitle("Form Layout")
        x = self.fill_data()
        if not x:
            pass
            #self.random_data_filler(40)
    def handleToolbar(self):
        self.toolbar = QToolBar()
        self.toolbar.setMovable(False)
        self.toolbar.setFloatable(False)
        self.toolbar.setToolButtonStyle(Qt.ToolButtonStyle.ToolButtonTextBesideIcon)
        self.toolbar.setStyleSheet("QToolBar { border: 0px; }")
        self.toolbar.setContentsMargins(0, 0, 0, 0)
        self.toolbar.setToolButtonStyle(Qt.ToolButtonStyle.ToolButtonTextBesideIcon)
        self.toolbar.setMovable(False)
        self.toolbar.setFloatable(False)
        self.toolbar.setOrientation(Qt.Orientation.Horizontal)#
        #make toolbar take up as little space as possible
        self.toolbar.setSizePolicy(QSizePolicy.Policy.Minimum, QSizePolicy.Policy.Minimum)

        self.toolbar.addAction("Import Data", self.handleImportData)
        self.toolbar.addAction("Save and Exit", self.handleSaveCurrent)
        return self.toolbar

    def handleImportData(self):
        #signal emit of (true, )
        #dialog selection of file

        x = QFileDialog.getOpenFileName(self, "Open File", "", "Excel Files (*.xlsx)")
        if x[0] != "":
            self.data = self.excel_handle_import(x[0])
            self.fill_data()
        
    def handleSaveCurrent(self):
        pass

    def createFormGroupBox(self):
        #create a main Vertical Layout
        
        self.formGroupBox = QGroupBox("Form layout")
        main_layout = QVBoxLayout()
        self.formGroupBox.setLayout(main_layout)

                #completed jobs goes in a QStackedWidget

        button_layout = QToolBar()
        main_layout.addWidget(button_layout)
        self.stackedWidget = QStackedWidget()
        main_layout.addWidget(self.stackedWidget)

        self.tables =  {
                "Completed Jobs":{"AV#":"int","W3W":"str","Location":"str","Sq/m Painted":"int","Sq/m Cleaned":"int"},
                "Inspections":{"W3W":"str","Location":"str","Notes":"str","TM":"bool"},
                "Inspection Routes":{"inspection route":"str","start":"str","end":"str","notes":"str"},
                "Schedule":{"W3W":"str","Location":"str","Notes":"str"},
                "Comments":{"Comments":"str"}
                }
        

        #create a vertical layout for completed_jobs
        #self.completedjobs = CompletedJobs(self.formGroupBox, self.session)
        #self.completed_jobs_button = QPushButton("Completed Jobs")
        #self.completed_jobs_button.clicked.connect(partial(self.stackedWidget.setCurrentIndex, 0))
        #self.stackedWidget.addWidget(self.completedjobs.widget)
        #button_layout.addWidget(self.completed_jobs_button)
        self.table_list = {}
        for k,v in self.tables.items():
            if k == "Completed Jobs":
                table = CompletedJobs(self.formGroupBox, self.session, v)
            else:
                table = TabSheets(self.formGroupBox, k, v)
            button = QPushButton(k)
            button.clicked.connect(partial(self.stackedWidget.setCurrentIndex, len(self.table_list)))
            self.stackedWidget.addWidget(table.widget)
            button_layout.addWidget(button)
            self.table_list[k] = table

    def closeEvent(self, a0) -> None:
        if not self._done:
            self.signal.emit((False, False))
        self.close()
        a0.accept()

    def accept(self):
        self._done = True
        self.emit_data()
        self.close()

    def emit_data(self):
        #completed_jobs_data = self.data_getter(self.completedjobs)
        all_data = {}
        for table_name, table in self.table_list.items():
            all_data[table_name] = self.data_getter(table)
        self.signal.emit((all_data, self.filename.text()))

    def data_getter(self, object_):
        #first item in object_data are column keys as an array
        object_data = [list(object_.columns.keys())]
        for row in range(object_.table.rowCount()):
            
            t_arr = []
            #add headers first
            #col1, 4, 5 = spinbox
            for col in range(object_.table.columnCount()):
                object_item = object_.table.cellWidget(row,col)
                if isinstance(object_item, QSpinBox):
                    t_arr.append(object_item.value())
                elif isinstance(object_item, QCheckBox):
                    t_arr.append(object_item.isChecked())
                else:
                    x = object_.table.item(row, col)
                    t_arr.append(x.text() if x else None)
            object_data.append(t_arr)
        return object_data

    def fill_data(self):
        #clear the table
        for table in self.table_list.values():
            table.table.setRowCount(0)

        if not self.data:
            return False
        #data len = 5, multidimensional array lists
        for key, value in self.data.items():
            if key in self.table_list:
                self.data_giver(self.table_list[key], value)

        return True

    def data_giver(self, object_, data):

        for row in range(len(data)):
            object_.add_btn.click()
            for col in range(len(data[row])):
                object_item = object_.table.cellWidget(row,col)
                #if object item is not qspinbox
                if type(object_item) == QSpinBox:
                    object_item.setValue(int(float(data[row][col])))
                elif type(object_item) == QCheckBox:
                    object_item.setChecked(bool(data[row][col]))
                elif not object_item:
                    #set item to be qtablewidgetitem
                    qtwi = QTableWidgetItem()
                    qtwi.setText(str(data[row][col]))

                    object_.table.setItem(row,col,qtwi)


