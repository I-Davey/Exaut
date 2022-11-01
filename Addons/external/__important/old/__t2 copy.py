from asyncio.log import logger
import inspect
from pydoc import doc
from re import S
from turtle import pd
#from __important.PluginInterface import PluginInterface
import os 
from PyQt6.QtWidgets import *
from PyQt6.QtCore import *
from functools import partial
import random
import time
from random_words  import RandomWords
import openpyxl
import pyperclip
import datetime

def excel_handle_import(file):
    #open excel file
    wb = openpyxl.load_workbook(file)

    #get all sheets
    sheets = wb.sheetnames

    #get all data from sheets
    data = {}
    for sheet in sheets:
        ws = wb[sheet]
        data[sheet] = []
        #all data except first row
        for row in ws.iter_rows(min_row=2):
            data[sheet].append([cell.value for cell in row])

    #return data
    return data






from docx import Document, table


#templace = C:\Users\idave\Documents\passbolt\t2\template.docx
#find all {{DATE}} and replace with current date
#find all tables and print them out

def replace_text(document, find, replace):
    for p in document.paragraphs:
        if find in p.text:
            inline = p.runs
            # Replace strings
            for i in range(len(inline)):
                if find in inline[i].text:
                    text = inline[i].text.replace(find, replace)
                    inline[i].text = text

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_text(cell, find, replace)

def get_tables(document):
    tables = []
    for table in document.tables:
        tables.append(table)
    return tables

def get_table_data(table):
    data = []
    for row in table.rows:
        data.append([])
        for cell in row.cells:
            data[-1].append(cell.text)
    return data

#row 2 col 1 has the table name
def get_table_name(table):
    return table.rows[1].cells[0].text

def add_table_data_noheader(table, data):
    #remove all but header row from table
    for i in range(len(table.rows)-1, 0, -1):
        table.rows[i].delete()
    #add data
    for row in data:
        table.add_row()
        for i in range(len(row)):
            table.rows[-1].cells[i].text = row[i]
print("hello world")
document = Document("C:\\Users\\idave\\Documents\\passbolt\\t2\\template.docx")
replace_text(document, "{{DATE}}", datetime.datetime.now().strftime("%d/%m/%Y"))
tables = get_tables(document)
for table_ in tables:
    print(get_table_name(table_))

